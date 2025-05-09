from flask import Flask, request, jsonify
from flask_cors import CORS
import fitz  # PyMuPDF
import os
import tempfile
import docx  # python-docx for Word documents
import json
from http.server import BaseHTTPRequestHandler
import base64
import io

app = Flask(__name__)
CORS(app)  # Enable CORS for all routes

def extract_text_from_pdf(pdf_content):
    text = ""
    try:
        # Create a file-like object from the PDF content
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp:
            temp.write(pdf_content)
            temp_path = temp.name
            
        # Open the PDF file
        doc = fitz.open(temp_path)
        
        # Iterate through each page
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text()
        
        # Close the document
        doc.close()
        
        # Clean up the temporary file
        os.unlink(temp_path)
        
    except Exception as e:
        text = f"Error extracting text from PDF: {str(e)}"
    
    return text

def extract_text_from_word(docx_content, is_docx=True):
    text = ""
    try:
        # Create a file-like object from the document content
        suffix = '.docx' if is_docx else '.doc'
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp:
            temp.write(docx_content)
            temp_path = temp.name
            
        # Open the Word document
        doc = docx.Document(temp_path)
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
            
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
            text += "\n"
        
        # Clean up the temporary file
        os.unlink(temp_path)
            
    except Exception as e:
        text = f"Error extracting text from Word document: {str(e)}"
    
    return text

def create_summary(text, file_name):
    """Create a summary of the extracted text"""
    # Count words
    words = len(text.split())
    
    # Get first few sentences (up to 3)
    sentences = text.replace('\n', ' ').split('.')
    first_sentences = '. '.join([s.strip() for s in sentences[:3] if s.strip()])
    
    summary = f"""
## Document Summary

**Document**: {file_name}  
**Word Count**: Approximately {words} words  
**Character Count**: {len(text)} characters

### Preview
{first_sentences}...

### Processing
This text was extracted using PyMuPDF for PDF files and python-docx for Word documents.
    """
    
    return summary

class handler(BaseHTTPRequestHandler):
    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()

    def do_POST(self):
        content_length = int(self.headers['Content-Length'])
        post_data = self.rfile.read(content_length)
        
        try:
            # Parse the JSON data
            data = json.loads(post_data)
            
            file_name = data.get('fileName', 'document')
            file_type = data.get('fileType', '').lower()
            file_content_base64 = data.get('fileContent', '')
            
            if not file_content_base64:
                self.send_error_response('No file content provided')
                return
                
            # Decode the base64 file content
            file_content = base64.b64decode(file_content_base64)
            
            # Extract text based on file type
            if '.pdf' in file_type or 'application/pdf' in file_type:
                text = extract_text_from_pdf(file_content)
            elif '.docx' in file_type or 'wordprocessingml' in file_type:
                text = extract_text_from_word(file_content, is_docx=True)
            elif '.doc' in file_type or 'msword' in file_type:
                text = extract_text_from_word(file_content, is_docx=False)
            else:
                self.send_error_response(f'Unsupported file type: {file_type}')
                return
                
            # Create a summary
            summary = create_summary(text, file_name)
            
            # Create the response
            response_data = {
                'text': text,
                'summary': summary,
                'fileName': file_name,
                'fileType': file_type
            }
            
            # Send the response
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            self.wfile.write(json.dumps(response_data).encode())
            
        except Exception as e:
            self.send_error_response(f'Error processing document: {str(e)}')
    
    def send_error_response(self, error_message):
        self.send_response(400)
        self.send_header('Content-type', 'application/json')
        self.send_header('Access-Control-Allow-Origin', '*')
        self.end_headers()
        self.wfile.write(json.dumps({'error': error_message}).encode()) 